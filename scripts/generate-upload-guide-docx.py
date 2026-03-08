"""Generate upload-guide.docx from the markdown content."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Style setup ---
style = doc.styles["Normal"]
style.font.name = "游ゴシック"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "游ゴシック"
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


# ============================================================
# Title
# ============================================================
title = doc.add_heading("データアップロード運用ガイド", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "本ドキュメントは、管理画面（admin）でのデータアップロード機能の仕様と運用手順をまとめたものです。"
)

# ============================================================
# 目次
# ============================================================
doc.add_heading("目次", level=1)
toc_items = [
    "1. 全体フロー",
    "2. 取引データインポート（MFクラウド CSV）",
    "3. 寄付者一括インポート（Donor CSV）",
    "4. ポートフォリオ資産インポート（Portfolio CSV）",
    "5. 定期更新の運用フロー",
    "6. トラブルシューティング",
    "7. テンプレートファイル一覧",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")

# ============================================================
# 全体フロー
# ============================================================
doc.add_heading("全体フロー", level=1)
doc.add_paragraph("データの投入から公開までの全体像は以下の通りです。")

flow_text = """MFクラウド会計
    ↓ CSVエクスポート
admin 管理画面
  1. 取引データインポート (/transactions)
     → MFクラウドCSVをアップロード
  2. 取引先（Counterpart）紐付け (/assign/counterparts)
     → 支出先の名前・住所を登録・紐付け
  3. 寄付者（Donor）紐付け (/assign/donors)
     → 個人寄附の寄付者情報を登録・紐付け
  4. ポートフォリオ資産登録 (/import-portfolio)
     → 期末の資産内訳をCSVで登録
  5. 報告書プロファイル設定 (/report-profile)
     → 団体名・住所等の報告書メタデータを設定
  6. 報告書XML出力 (/export-report)
     → 政治資金収支報告書XMLをダウンロード
    ↓ 自動反映
webapp 公開サイト
  → 収支グラフ・サンキー図・取引一覧が更新"""

p = doc.add_paragraph()
run = p.add_run(flow_text)
run.font.name = "Consolas"
run.font.size = Pt(9)

# ============================================================
# 1. 取引データインポート
# ============================================================
doc.add_heading("1. 取引データインポート（MFクラウド CSV）", level=1)

doc.add_heading("概要", level=2)
doc.add_paragraph(
    "MFクラウド会計からエクスポートした仕訳帳CSVをアップロードし、取引データとしてシステムに取り込みます。"
)
doc.add_paragraph("管理画面パス: /transactions（取引一覧ページのCSVアップロードフォーム）")
doc.add_paragraph("テンプレート: docs/templates/mf-transactions-template.csv")

doc.add_heading("CSV仕様", level=2)
add_table(
    doc,
    ["ヘッダー名", "別名", "必須", "説明"],
    [
        ["取引No", "—", "Yes", "取引の一意識別番号"],
        ["取引日", "—", "Yes", "取引日（YYYY-MM-DD 形式）"],
        ["借方勘定科目", "—", "Yes", "借方の勘定科目名"],
        ["借方補助科目", "—", "No", "借方の補助科目名"],
        ["借方部門", "—", "No", "借方の部門名"],
        ["借方取引先", "—", "No", "借方の取引先名"],
        ["借方税区分", "借方税区", "No", "借方の税区分"],
        ["借方インボイス", "—", "No", "借方のインボイス情報"],
        ["借方金額", "借方金額(円)", "Yes", "借方の金額（整数、カンマ区切り可）"],
        ["貸方勘定科目", "—", "Yes", "貸方の勘定科目名"],
        ["貸方補助科目", "—", "No", "貸方の補助科目名"],
        ["貸方部門", "—", "No", "貸方の部門名"],
        ["貸方取引先", "—", "No", "貸方の取引先名"],
        ["貸方税区分", "貸方税区", "No", "貸方の税区分"],
        ["貸方インボイス", "—", "No", "貸方のインボイス情報"],
        ["貸方金額", "貸方金額(円)", "Yes", "貸方の金額（整数、カンマ区切り可）"],
        ["摘要", "—", "No", "取引の摘要（説明文）"],
        ["タグ", "起訖タグ", "No", "表示用カテゴリ（friendlyCategory）"],
        ["メモ", "—", "No", "メモ"],
    ],
)

doc.add_heading("対応勘定科目一覧", level=2)

doc.add_heading("収入科目（PL - 貸方）", level=3)
add_table(
    doc,
    ["勘定科目名", "category_key", "報告書上の分類"],
    [
        ["個人の負担する党費又は会費", "membership-fees", "機関紙誌+その他事業収入"],
        ["個人からの寄附", "individual-donations", "寄附"],
        ["個人からの寄附（特定寄附）", "specific-individual-donations", "寄附"],
        ["法人その他の団体からの寄附", "corporate-donations", "寄附"],
        ["政治団体からの寄附", "political-donations", "寄附"],
        ["政党匿名寄附", "anonymous-donations", "寄附"],
        ["機関紙誌の発行その他の事業による収入", "publication-income", "機関紙誌+その他事業収入"],
        ["借入金", "loans", "借入金"],
        ["本部又は支部から供与された交付金に係る収入", "grants", "交付金"],
        ["政治資金パーティーの対価に係る収入", "party-income", "パーティー収入"],
        ["寄附のあっせんによるもの", "mediated-donations", "寄附"],
        ["政治資金パーティー対価のあっせんによるもの", "mediated-party-income", "パーティー収入"],
        ["その他の収入", "other-income", "その他"],
    ],
)

doc.add_heading("支出科目（PL - 借方）", level=3)
add_table(
    doc,
    ["勘定科目名", "category_key", "報告書上の分類"],
    [
        ["人件費", "personnel-costs", "経常経費"],
        ["光熱水費", "utilities", "経常経費"],
        ["備品・消耗品費", "equipment-supplies", "経常経費"],
        ["事務所費", "office-expenses", "経常経費"],
        ["組織活動費", "organizational-activities", "政治活動費"],
        ["選挙関係費", "election-expenses", "政治活動費"],
        ["機関紙誌の発行事業費", "publication-expenses", "政治活動費"],
        ["宣伝事業費", "advertising-expenses", "政治活動費"],
        ["政治資金パーティー開催事業費", "fundraising-party-expenses", "政治活動費"],
        ["その他の事業費", "other-business-expenses", "政治活動費"],
        ["調査研究費", "research-expenses", "政治活動費"],
        ["寄附・交付金", "donations-grants-expenses", "政治活動費"],
        ["その他の経費", "other-expenses", "政治活動費"],
    ],
)

doc.add_heading("貸借対照表科目（BS）", level=3)
add_table(
    doc,
    ["勘定科目名", "種別"],
    [
        ["普通預金", "資産（現金類）"],
        ["未払金/未払費用", "負債"],
    ],
)

doc.add_heading("特殊科目", level=3)
add_table(
    doc,
    ["勘定科目名", "説明"],
    [
        ["相殺項目（費用）", "相殺支出仕訳"],
        ["相殺項目（収入）", "相殺収入仕訳"],
    ],
)

doc.add_heading("取引種別の自動判定ルール", level=2)
add_table(
    doc,
    ["条件", "判定結果"],
    [
        ["借方 = 相殺項目（費用）", "offset_expense"],
        ["貸方 = 相殺項目（収入）", "offset_income"],
        ["借方 = 現金類（BS） & 貸方 = PL科目", "income（現金収入）"],
        ["借方 = PL科目 & 貸方 = 現金類（BS）", "expense（現金支出）"],
        ["借方 = PL & 貸方 = BS（現金以外）、またはその逆", "non_cash_journal（非現金仕訳）"],
        ["上記いずれにも該当しない", "invalid（エラー）"],
    ],
)

doc.add_heading("アップロード手順", level=2)
steps = [
    "MFクラウド会計の「仕訳帳」からCSVをエクスポート",
    "admin管理画面 /transactions を開く",
    "政治団体を選択",
    "CSVファイルを選択",
    "プレビュー画面で以下を確認:\n  - 新規追加（insert）件数\n  - 更新（update）件数\n  - スキップ（skip）件数 — 既にハッシュ値が同一のデータ\n  - エラー（invalid）件数 — 勘定科目の不一致など",
    "問題がなければ「このデータを保存する」をクリック",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("重複検知の仕組み", level=2)
doc.add_paragraph("取引No と 政治団体ID の組み合わせで既存データを照合します。")
add_table(
    doc,
    ["条件", "ステータス", "動作"],
    [
        ["取引Noが存在しない", "insert", "新規追加"],
        ["取引Noが存在 & ハッシュ同一", "skip", "スキップ（重複防止）"],
        ["取引Noが存在 & ハッシュ異なる", "update", "データ更新"],
    ],
)

doc.add_heading("エンコーディング", level=2)
doc.add_paragraph(
    "Shift_JIS と UTF-8 の両方に対応しています。MFクラウドからのエクスポートCSV（通常Shift_JIS）をそのままアップロードできます。"
)

# ============================================================
# 2. 寄付者一括インポート
# ============================================================
doc.add_heading("2. 寄付者一括インポート（Donor CSV）", level=1)

doc.add_heading("概要", level=2)
doc.add_paragraph(
    "個人寄附の取引に対して、寄付者情報（名前・住所・職業等）を一括で紐付けるためのCSVインポート機能です。"
    "政治資金収支報告書では、5万円を超える個人寄附について寄付者情報の記載が必要です。"
)
doc.add_paragraph("管理画面パス: /import-donors")
doc.add_paragraph("テンプレート: docs/templates/donor-template.csv")
doc.add_paragraph("最大行数: 1,000行")

doc.add_heading("CSV仕様", level=2)
add_table(
    doc,
    ["ヘッダー名", "必須", "説明"],
    [
        ["取引No", "Yes", "紐付け対象の取引番号"],
        ["寄付者名", "Yes", "寄付者のフルネーム（最大120文字）"],
        ["寄付者種別", "Yes", "individual / corporation / political_organization"],
        ["住所", "Yes", "寄付者の住所（最大120文字、空欄可）"],
        ["職業", "Yes", "寄付者の職業（最大50文字、空欄可）"],
    ],
)

doc.add_heading("寄付者種別一覧", level=2)
add_table(
    doc,
    ["値", "意味"],
    [
        ["individual", "個人"],
        ["corporation", "法人その他の団体"],
        ["political_organization", "政治団体"],
    ],
)

doc.add_heading("アップロード手順", level=2)
for i, step in enumerate(
    [
        "テンプレートCSVをダウンロードし、寄付者情報を記入",
        "admin管理画面 /import-donors を開く",
        "政治団体を選択",
        "CSVファイルを選択してプレビュー",
        "プレビュー画面で紐付け先の取引情報と寄付者情報を確認",
        "問題なければ「インポート」をクリック",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("注意事項", level=2)
for note in [
    "取引No は既にインポート済みの取引データに存在する番号を指定してください",
    "同一取引に対して複数の寄付者を紐付けることはできません（1取引 = 1寄付者）",
    "既に寄付者が紐付いている取引に対して再度インポートすると上書きされます",
]:
    doc.add_paragraph(note, style="List Bullet")

# ============================================================
# 3. ポートフォリオ資産インポート
# ============================================================
doc.add_heading("3. ポートフォリオ資産インポート（Portfolio CSV）", level=1)

doc.add_heading("概要", level=2)
doc.add_paragraph(
    "期末時点での資産内訳（現金・株式・貴金属・不動産等）を登録するためのCSVインポート機能です。"
    "webapp公開サイトのポートフォリオ表示に反映されます。"
)
doc.add_paragraph("管理画面パス: /import-portfolio")
doc.add_paragraph("テンプレート: docs/templates/portfolio-template.csv")
doc.add_paragraph("最大行数: 100行")

doc.add_heading("CSV仕様", level=2)
add_table(
    doc,
    ["ヘッダー名", "必須", "説明"],
    [
        ["category", "Yes", "資産カテゴリ（下記参照）"],
        ["label", "Yes", "資産の表示名（例: 普通預金（みずほ銀行））"],
        ["amount", "Yes", "金額（正の整数）"],
        ["snapshotDate", "Yes", "基準日（YYYY-MM-DD 形式）"],
    ],
)

doc.add_heading("資産カテゴリ一覧", level=2)
add_table(
    doc,
    ["category値", "意味"],
    [
        ["cash", "現金・預金"],
        ["stocks", "有価証券（株式・投資信託等）"],
        ["precious_metals", "貴金属"],
        ["real_estate", "不動産"],
        ["other", "その他の資産"],
    ],
)

doc.add_heading("アップロード手順", level=2)
for i, step in enumerate(
    [
        "テンプレートCSVをダウンロードし、資産情報を記入",
        "admin管理画面 /import-portfolio を開く",
        "政治団体を選択",
        "CSVファイルの内容をテキストエリアに貼り付けるか、ファイルを選択",
        "「インポート」をクリック",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("注意事項", level=2)
for note in [
    "同日付のデータは上書きされます: 同じ snapshotDate のデータが既に存在する場合、その日付の全データが削除されてから新しいデータが挿入されます",
    "定期的に更新する場合は、対象日付のデータを全行含めたCSVを用意してください（差分更新ではなく全件置換）",
    "amount は0以上の整数を指定してください",
]:
    doc.add_paragraph(note, style="List Bullet")

# ============================================================
# 定期更新の運用フロー
# ============================================================
doc.add_heading("定期更新の運用フロー", level=1)

doc.add_heading("月次更新（推奨）", level=2)
doc.add_paragraph("毎月の会計データを反映するための定期運用手順です。")

monthly_steps = [
    ("Step 1", "MFクラウドから仕訳帳CSVをエクスポート", "対象期間: 当該年度の1月1日〜最終月末日。全期間分をエクスポートしてOK（重複は自動スキップ）。"),
    ("Step 2", "admin /transactions でCSVアップロード", "プレビューで insert/update/skip 件数を確認。"),
    ("Step 3", "新しい取引に対して取引先を紐付け", "admin /assign/counterparts。AI候補機能で住所・郵便番号を自動取得可能。"),
    ("Step 4", "寄附取引に寄付者情報を紐付け（該当がある場合）", "admin /assign/donors または /import-donors（一括）。"),
    ("Step 5", "webapp公開サイトで反映を確認", ""),
]
for step_name, desc, note in monthly_steps:
    p = doc.add_paragraph()
    run = p.add_run(f"{step_name}: {desc}")
    run.bold = True
    if note:
        doc.add_paragraph(f"  {note}")

doc.add_heading("年次更新（決算期末）", level=2)
doc.add_paragraph("年度末に追加で必要な作業です。")

yearly_steps = [
    ("Step 1", "月次更新の Step 1〜4 を実施（12月分まで）"),
    ("Step 2", "ポートフォリオ資産CSVを作成し、admin /import-portfolio でアップロード（期末日: 例 2025-12-31）"),
    ("Step 3", "残高スナップショットを登録（admin /balance-snapshots）"),
    ("Step 4", "報告書プロファイルを確認・更新（admin /political-organizations/[orgId]/report-profile）"),
    ("Step 5", "報告書プレビューで内容を確認（admin /export-report/[orgId]/[year]）"),
    ("Step 6", "政治資金収支報告書XMLをダウンロード"),
    ("Step 7", "総務省・選挙管理委員会に提出"),
]
for step_name, desc in yearly_steps:
    p = doc.add_paragraph()
    run = p.add_run(f"{step_name}: {desc}")
    run.bold = True

doc.add_heading("MFクラウドからのCSVエクスポート手順", level=2)
for i, step in enumerate(
    [
        "MFクラウド会計にログイン",
        "「仕訳帳」メニューを開く",
        "対象期間を設定（例: 2025/01/01 〜 2025/12/31）",
        "「エクスポート」→「CSV出力」を選択",
        "ダウンロードしたCSVファイルをそのまま admin にアップロード",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}")

p = doc.add_paragraph()
run = p.add_run("ポイント: ")
run.bold = True
p.add_run(
    "MFクラウドのCSVはShift_JIS形式ですが、システムが自動判定してデコードします。"
    "エクセル等で開いて再保存すると文字化けの原因になるため、ダウンロードしたファイルはそのままアップロードしてください。"
)

# ============================================================
# トラブルシューティング
# ============================================================
doc.add_heading("トラブルシューティング", level=1)

troubles = [
    (
        "「無効な借方科目」「無効な貸方科目」エラー",
        "MFクラウドの勘定科目名がシステムで認識されていません。",
        [
            "MFクラウドの勘定科目名が対応勘定科目一覧に含まれているか確認",
            "科目名が完全一致していない場合（全角/半角の違いなど）はMFクラウド側で修正してから再エクスポート",
        ],
    ),
    (
        "「独自のカテゴリが設定されていません」エラー",
        "MFクラウドの「タグ（起訖タグ）」列が空欄の取引があります。",
        [
            "MFクラウドでタグ（起訖タグ）が設定されていない仕訳にタグを設定してから再エクスポート",
            "相殺項目や非現金仕訳はタグ不要です",
        ],
    ),
    (
        "CSVが文字化けする",
        "",
        [
            "MFクラウドからダウンロードしたCSVを、Excel等で開かずにそのままアップロード",
            "Excel等で編集が必要な場合は、保存時に文字コードをUTF-8（BOMなし）にする",
        ],
    ),
    (
        "全件スキップされる",
        "既にアップロード済みのデータと同一内容です（ハッシュ値が一致）。",
        [
            "データに変更がないことの確認なので、正常な動作です",
            "MFクラウドで修正した取引がある場合は、修正後にCSVを再エクスポートしてアップロード",
        ],
    ),
    (
        "ポートフォリオのデータが消えた",
        "同日付の既存データは全件削除されてから新規挿入されます。",
        [
            "同じ snapshotDate の資産情報は、必ず全行をCSVに含めてアップロードしてください",
            "差分更新はできないため、部分的なCSVをアップロードすると既存データが失われます",
        ],
    ),
]

for title, desc, fixes in troubles:
    doc.add_heading(title, level=2)
    if desc:
        doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run("対処法:")
    run.bold = True
    for fix in fixes:
        doc.add_paragraph(fix, style="List Bullet")

# ============================================================
# テンプレートファイル一覧
# ============================================================
doc.add_heading("テンプレートファイル一覧", level=1)
add_table(
    doc,
    ["テンプレート", "ファイルパス", "用途"],
    [
        ["取引データ（MFクラウド形式）", "docs/templates/mf-transactions-template.csv", "MFクラウドCSVインポート"],
        ["寄付者一括インポート", "docs/templates/donor-template.csv", "寄付者情報の一括紐付け"],
        ["ポートフォリオ資産", "docs/templates/portfolio-template.csv", "期末資産内訳の登録"],
    ],
)

doc.add_paragraph(
    "テンプレートファイルにはサンプルデータが含まれています。ヘッダー行はそのまま使用し、データ行を実際の値に置き換えてください。"
)

# ============================================================
# Save
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "..", "docs", "upload-guide.docx")
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f"Generated: {output_path}")
